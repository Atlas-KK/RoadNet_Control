from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables" / "陕西交控路网安全态势感知大屏UI设计文档.docx"
IMAGES = {
    "overview": Path(r"C:\Users\12040\AppData\Local\Temp\codex-clipboard-3aed5c86-ee20-4137-9d54-1ad7c867829c.png"),
    "event": Path(r"C:\Users\12040\AppData\Local\Temp\codex-clipboard-c931f9bf-2017-4b02-b0d5-b68a46bb87f4.png"),
    "maintenance": Path(r"C:\Users\12040\AppData\Local\Temp\codex-clipboard-7d9446bb-7f37-4727-b754-2e0f2f63429e.png"),
}

NAVY = "0B1F33"
INK = "15324D"
BLUE = "1E74D8"
CYAN = "00B7F4"
TEAL = "00CFD5"
LIGHT = "EAF3FB"
MUTED = "5D7184"
GRID = "B8C9D8"
PALE = "F4F8FC"
AMBER = "F5A623"
RED = "D84A52"
GREEN = "24B47E"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "bottom", "start", "end"):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "80" if side in ("top", "bottom") else "120")
                node.set(qn("w:type"), "dxa")


def border_bottom(paragraph, color=BLUE, size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_run(run, size=10.5, color=INK, bold=False, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_para(doc, text="", size=10.5, color=INK, bold=False, after=6, before=0, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size, color, bold)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        set_run(p.add_run(item), 10.3)


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(text), 16, BLUE, True)
    border_bottom(p, CYAN, "8")
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(text), 12.5, BLUE, True)
    return p


def table(doc, headers, rows, widths, header_fill="E8EEF5", font_size=9.2):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    for i, value in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, header_fill)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(value), font_size, "173A5E", True)
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            if r_idx % 2 == 1:
                shade(cells[i], "F7FAFD")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.08
            set_run(p.add_run(value), font_size, INK)
    set_table_geometry(t, widths)
    return t


def add_callout(doc, title, text, color=BLUE):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(t, [9360])
    c = t.cell(0, 0)
    shade(c, "EFF7FD")
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title + "  "), 10.5, color, True)
    set_run(p.add_run(text), 10.3, INK)
    return t


def add_caption(doc, text):
    p = add_para(doc, text, size=8.5, color=MUTED, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    return p


def add_image(doc, image, width, caption):
    if image.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(image), width=Inches(width))
        add_caption(doc, caption)


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.68)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)
    sec.header_distance = Inches(0.32)
    sec.footer_distance = Inches(0.30)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run(header.add_run("陕西交控路网安全态势感知 | UI设计文档"), 8.5, MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("陕西交控路网安全态势感知大屏 · 设计基线 9568 × 2808"), 8.5, MUTED)


def page_break(doc):
    doc.add_page_break()


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    setup(doc)

    # Cover
    add_para(doc, "SHAANXI JIAOKONG", 10, CYAN, True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "陕西交控路网安全态势感知", 27, NAVY, True, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "超宽大屏 UI 设计文档", 18, BLUE, True, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = add_para(doc, "基于 G70 福银高速西长段运行监测辅助决策驾驶舱参考界面梳理", 10.5, MUTED, after=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    border_bottom(p, CYAN, "14")
    table(doc, ["项目", "设计基线", "文档定位"], [["陕西交控路网安全态势感知大屏", "9568 × 2808 px（约 3.41:1）", "产品、视觉、前端与数据接入的共同设计依据"]], [2800, 2750, 3810], "EAF5FC", 9.5)
    add_para(doc, "版本：V1.0    日期：2026-08-07    状态：方案梳理", 9.5, MUTED, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_image(doc, IMAGES["overview"], 6.75, "图 1 参考概览页：中心地图 + 左侧态势总览 + 右侧专题详情的三域结构")
    add_callout(doc, "设计结论", "该大屏以‘一张路网图’承载态势定位，以两翼抽屉承载安全研判，以顶部指标带承载资源总览；通过右侧专题切换覆盖交通、气象、结构物、设备等安全子域。", CYAN)

    page_break(doc)
    h1(doc, "1. 设计目标与适用范围")
    add_para(doc, "本方案面向陕西交控路网运行监测与安全指挥场景，支持超宽拼接屏常态值守、事件研判、养护协同和应急调度。参考界面保留深色数字孪生地图、蓝青科技光效、左中右三域信息架构与专题大屏扩展方式。")
    h2(doc, "1.1 核心目标")
    add_bullets(doc, [
        "30 秒内识别当前路网安全等级、事件数量、交通管制、气象预警和关键结构物异常。",
        "在中心地图完成路段、设施、事件、视频和设备的空间定位、联动筛选与下钻。",
        "在不离开主地图的前提下，通过右侧专题面板切换完成风险研判与处置核验。",
        "兼容 9568 × 2808 的拼接屏，关键数字、文字和点击目标均按远距观看与值守操作设计。",
    ])
    h2(doc, "1.2 用户与任务")
    table(doc, ["角色", "主要任务", "优先信息"], [
        ["路网值守人员", "监控、确认、派单、跟踪", "实时事件、视频、拥堵、设备故障"],
        ["安全管理人员", "风险研判、会商、督导", "综合风险等级、预警、结构物、管制影响"],
        ["养护与应急人员", "施工协同、巡查、处置", "施工计划、桥隧异常、值班资源、应急物资"],
        ["管理决策人员", "态势总览、趋势复盘", "事件趋势、通行效率、处置时效、风险闭环"],
    ], [1700, 3600, 4060])
    h2(doc, "1.3 页面体系")
    table(doc, ["层级", "页面 / 模式", "功能定位"], [
        ["L0", "综合概览", "一图统览当前安全态势，作为默认值守页。"],
        ["L1", "右侧专题抽屉", "通行、管制、气象、关键结构物、设备、事件明细等专题切换。"],
        ["L1", "交通流及事件", "趋势、流量、收费站排行、断面与车型分析。"],
        ["L1", "养护巡查及运营", "施工统计、关键结构物卡片、值班值守、路产巡查。"],
        ["L2", "对象详情", "事件、设施、设备、收费站、桥隧的详情、视频与处置时间线。"],
    ], [900, 2400, 6060])

    page_break(doc)
    h1(doc, "2. 总体布局与分辨率基线")
    add_para(doc, "画布固定为 9568 × 2808 px，建议按逻辑栅格设计、按屏幕比例缩放。业务区域不使用浏览器滚动；右侧专题内容区允许卡片内纵向滚动。拼接缝应避开标题、核心数值、图表零轴与高频按钮。")
    h2(doc, "2.1 布局坐标建议")
    table(doc, ["区域", "逻辑坐标（x, y, w, h）", "占比", "设计说明"], [
        ["顶栏", "0, 0, 9568, 220", "7.8% 高", "系统时间、主标题、页面模式选择；标题居中，控制项靠右。"],
        ["左侧态势栏", "90, 250, 2240, 2490", "23.4% 宽", "固定 6 个信息组：事件、通行、管制、气象、结构物、设备。"],
        ["中心地图", "2360, 250, 4760, 2490", "49.8% 宽", "主路网空间底图、上方设施 KPI、右下图层控制。"],
        ["右侧专题栏", "7150, 250, 2328, 2490", "24.3% 宽", "同一容器随专题切换内容，承载研判、排行、趋势与明细。"],
        ["安全边距", "≥ 90 px", "—", "所有贴边模块保留外呼吸边距，适配边缘拼缝与投屏裁切。"],
    ], [1300, 3000, 1050, 4010])
    h2(doc, "2.2 响应式与缩放规则")
    add_bullets(doc, [
        "设计稿以 9568 px 为 100% 基准；现场使用 scale = 实际屏宽 / 9568 等比缩放，字体、图标、圆角、阴影同步缩放。",
        "最小字号：非核心说明 28 px；表头与列表正文 30–34 px；模块标题 38–44 px；关键指标 64–88 px；主标题 78–92 px。",
        "左右栏宽度为固定比例，中心地图为弹性区；当可用宽度下降时优先压缩地图，不压缩左右栏的文字与表格列宽。",
        "功能操作区的最小可点击尺寸为 72 × 48 px；地图图例与勾选框保持 24 px 以上视觉尺寸。",
    ])
    add_callout(doc, "拼接屏规范", "标题、KPI、图表标签和关键告警不得横跨物理拼缝。建议在实施前录入实际屏体分辨率与拼缝坐标，生成拼缝避让参考线。", AMBER)

    page_break(doc)
    h1(doc, "3. 概览页信息架构")
    add_para(doc, "概览页采用‘顶部资源总览 + 左侧安全态势 + 中心一图定位 + 右侧专题研判’架构。左栏提供连续摘要，中心地图提供空间上下文，右栏展示当前选中专题，不抢占主地图。")
    h2(doc, "3.1 顶部资源总览")
    table(doc, ["组件", "字段", "状态 / 交互"], [
        ["系统时间", "日期、时分秒、数据更新时间", "每秒刷新时间；数据更新时间独立标识。"],
        ["主标题", "陕西交控路网安全态势感知", "固定居中；可选显示当前机构 / 路段范围。"],
        ["页面模式", "概况、交通流及事件、养护巡查及运营", "下拉或分段按钮切换一级页面。"],
        ["设施 KPI", "总里程、收费站、服务区、桥梁、隧道", "点击筛选地图并打开对应专题或图层。"],
    ], [1600, 3600, 4160])
    h2(doc, "3.2 左侧固定态势栏")
    table(doc, ["模块", "建议指标", "告警表现"], [
        ["交通事件", "实时事件总数；人员伤亡、路产受损、阻断交通、道路施工", "总数使用高亮数字；阻断交通使用红色；点击定位事件。"],
        ["通行态势", "拥堵等级、拥堵路段数、重大项目影响长度", "绿/黄/橙/红四级；显示较昨日或近 15 分钟变化。"],
        ["交通管制", "封道里程、关闭、限流、分流", "管制类型与影响长度并列；与地图线段高亮联动。"],
        ["气象环境", "气象预警、大风、降雨、降雪、大雾、其他", "按预警级别着色；点击转到气象预警列表。"],
        ["关键结构物", "桥梁异常、边坡异常、隧道异常 / 总数", "异常分子显红，分母为纳管对象数；可横向切换。"],
        ["机电设备", "完好率、设备总数、故障数", "环图显示完好率；故障数红色并提供明细入口。"],
    ], [1700, 4250, 3550])
    h2(doc, "3.3 中心地图与图层")
    add_bullets(doc, [
        "底图采用深海军蓝底色，主路网蓝灰线、河流青色、重点对象青绿标注；低优先级地物降亮度，确保事件与管制线段显著。",
        "地图上方悬浮设施 KPI，避免覆盖路网核心区；图例固定在右下，支持多选图层与全选/清空。",
        "图层至少包含：实时事件、道路施工、交通管制、道路拥堵、巡查车辆、应急资源、服务区、收费站、桥梁、隧道、摄像机、情报板。",
        "点选对象弹出轻量浮卡：名称、桩号/位置、风险等级、最近上报时间、操作按钮（查看详情 / 视频 / 派单）。",
    ])

    page_break(doc)
    h1(doc, "4. 右侧专题抽屉设计")
    add_para(doc, "右侧专题区域统一使用 2328 px 宽的垂直容器。标题行固定，主体根据专题切换；内容卡片使用 24 px 内边距、12–16 px 圆角、1 px 青蓝描边与低透明深蓝渐变填充。")
    table(doc, ["专题", "参考图对应内容", "推荐组件与交互"], [
        ["实时事件", "近 15 日趋势、影响统计、类型分析、事件明细", "折线 / 柱线组合、环形占比、状态 Tab、空态插画；点击行定位地图。"],
        ["通行概况", "主线与收费站切换、拥堵清单、今日流量、24 小时趋势", "分段按钮、表格、柱线组合图；筛选收费站与入口/出口。"],
        ["交通管制", "主线/收费站管制、原因分析、管制详情", "KPI 条、同心环、详情卡；地图显示封控起讫与影响范围。"],
        ["气象环境", "蓝/黄/橙/红预警、预警卡片列表", "预警等级卡、滚动消息卡、天气预报 Tab；点击卡片高亮影响路段。"],
        ["关键结构物", "桥/边坡/隧道异常、对象列表", "摘要条 + 对象卡；包含传感器在线、健康状态、视频入口。"],
        ["设备概况", "设备总数、故障数、分类统计、故障明细", "设备类别 Tab、完好率、明细表；支持工单状态筛选。"],
    ], [1400, 3480, 4620], "EAF5FC", 8.9)
    h2(doc, "4.1 专题通用状态")
    table(doc, ["状态", "视觉规则", "交互 / 文案"], [
        ["默认", "选中 Tab 采用亮蓝填充，未选中使用透明描边。", "显示最近刷新时间与数据范围。"],
        ["告警", "红色用于阻断、故障、一级预警；橙色用于较高风险；黄色用于注意。", "高风险卡片置顶，提供‘定位’和‘查看处置’。"],
        ["加载", "保留卡片骨架，避免布局跳动。", "显示‘数据加载中’，超过阈值可重试。"],
        ["空态", "使用线性图标与浅色说明，不显示空图表。", "示例：‘当前筛选条件下暂无未处置事件’。"],
        ["异常", "模块级错误提示，不遮挡地图与其他模块。", "示例：‘设备数据暂不可用，已保留最近一次有效结果’。"],
    ], [1150, 4100, 4250])

    page_break(doc)
    h1(doc, "5. 专题全屏页面")
    h2(doc, "5.1 交通流及事件")
    add_para(doc, "参考图采用三栏分析布局：左栏关注事件量与时效，中栏关注出入口流量与收费站排名，右栏关注断面流量与车型结构。整体适合用于事件复盘、早晚高峰研判和收费站运行分析。")
    table(doc, ["区域", "模块", "关键字段"], [
        ["左栏", "事件数量变化、事件类型分析、近 12 月时效分析、路段排行", "今日/本周/本月/本年事件数、同比环比、上报/确认/处置/总计时效"],
        ["中栏", "出入口流量、近 13 月统计、收费站流量排行", "客车/货车/总计、入口/出口、本年本月本日、收费站名称与分类流量"],
        ["右栏", "近 13 月断面流量、断面车型占比", "断面名称、车型、入口/出口、占比滑杆与总量"],
    ], [1100, 3400, 5000])
    add_image(doc, IMAGES["event"], 6.9, "图 2 交通流及事件专题：以趋势图、排行表和车型占比支持复盘分析")
    h2(doc, "5.2 养护巡查及运营")
    add_para(doc, "参考图采用三栏运营布局：左侧统计施工规模与计划，中心滚动呈现关键结构物和视频，右侧展示值班人员与路产巡查。此页强化‘施工—设施—人员’的协同关系。")
    table(doc, ["区域", "模块", "关键字段 / 操作"], [
        ["左栏", "施工数量、施工位置、未来 30 日施工", "临时养护 / 中大修，主线/桥梁/隧道/收费站/服务区分布，日期趋势"],
        ["中栏", "关键结构物列表", "名称、类型、长度、天气、实时车流、上次告警、视频缩略图、告警状态"],
        ["右栏", "值班值守、路产巡查", "值班领导 / 人员 / 电话、养护中心、巡查班组、班次、车辆、人员"],
    ], [1100, 3400, 5000])
    add_image(doc, IMAGES["maintenance"], 6.9, "图 3 养护巡查及运营专题：以施工、结构物、值班巡查形成运营闭环")

    page_break(doc)
    h1(doc, "6. 视觉语言与组件规范")
    h2(doc, "6.1 颜色与风险映射")
    table(doc, ["角色", "色值建议", "使用边界"], [
        ["空间底色", "#061526 / #0B1F33", "页面底色、地图底色；禁止大面积纯黑。"],
        ["主科技蓝", "#1E74D8", "选中态、关键线条、柱状图主系列、一级操作。"],
        ["高亮青", "#00B7F4 / #00CFD5", "地图高亮、流量、正常在线、光效点缀。"],
        ["正常", "#24B47E", "正常、已恢复、健康状态；不得与高亮青混用表达风险。"],
        ["注意", "#F5C542", "轻度拥堵、黄色预警、待关注。"],
        ["较高风险", "#F28C38", "中度风险、施工影响、橙色预警。"],
        ["高风险", "#D84A52", "阻断、故障、红色预警、未处置关键事件。"],
    ], [1600, 2300, 5600])
    h2(doc, "6.2 字体、间距与图表")
    add_bullets(doc, [
        "中文优先使用 Microsoft YaHei / 思源黑体，数字使用 DIN Condensed 或等宽数字字体；数值与单位分层，核心数字加粗。",
        "模块标题采用‘双斜杠 + 标题 + 下划线光带’；标题区高度建议 96–112 px，正文区采用 32 px 基准间距。",
        "图表背景透明；网格线低对比度；最多使用 4 个业务系列，固定颜色映射并在所有专题保持一致。",
        "环图仅用于占比和健康度；当分类超过 5 项时改为横向条形图或排序列表，避免图例拥挤。",
        "表格行高建议 ≥ 64 px；表头 64–72 px；斑马纹透明度低于 12%；长位置字段允许截断并提供悬浮全量文本。",
    ])
    h2(doc, "6.3 通用组件规格")
    table(doc, ["组件", "尺寸 / 形态", "使用规则"], [
        ["模块容器", "描边 1 px；圆角 12–16 px；内边距 24–32 px", "采用深蓝半透明渐变，阴影只用于焦点层级。"],
        ["KPI 数字", "数字 64–88 px；单位 28–34 px", "数字与单位基线对齐；风险数字使用语义色。"],
        ["Tab / 筛选", "高度 48–56 px；水平间距 12–16 px", "选中有填充且可识别，不只依赖颜色。"],
        ["地图标注", "主标注 30–34 px；点位图标 44–56 px", "相邻标注自动避让；同屏只突出高优先级对象。"],
        ["视频缩略图", "16:9；最小 260 × 146 px", "无流时显示摄像机状态与重连操作，不显示纯黑占位。"],
    ], [1600, 3200, 4700])

    page_break(doc)
    h1(doc, "7. 交互、数据与刷新策略")
    h2(doc, "7.1 关键交互链路")
    table(doc, ["触发", "系统反馈", "后续动作"], [
        ["点击左栏风险指标", "地图按对象类型过滤并高亮，右栏切换对应专题", "点击列表项打开对象浮卡或详情。"],
        ["点击顶部设施 KPI", "地图聚焦该设施类型并显示数量 / 异常摘要", "支持继续筛选桥、隧、收费站、服务区。"],
        ["点击地图事件点", "显示事件浮卡与影响范围，右栏定位明细行", "查看视频、处置过程、关联管制与附近资源。"],
        ["切换右栏 Tab", "保留当前地图视野与图层，更新右栏数据", "不刷新全页，避免值守视线中断。"],
        ["选择时间 / 路段 / 收费站", "所有相关图表显示统一口径和加载状态", "筛选条件在专题顶部可见，可一键清除。"],
    ], [1800, 4000, 3700])
    h2(doc, "7.2 数据刷新与可信度")
    table(doc, ["数据类型", "建议刷新", "表现方式"], [
        ["实时事件、管制、设备告警", "15–30 秒或消息推送", "模块显示‘数据截至 HH:mm:ss’；新入事件弱闪烁一次后静止。"],
        ["交通流、断面、收费站", "1–5 分钟", "曲线平滑更新，显示数据周期与统计口径。"],
        ["气象预警", "5 分钟 + 预警推送", "有预警立即更新卡片；过期预警自动置灰并归档。"],
        ["桥隧监测与健康度", "1–5 分钟 / 阈值触发", "显示传感器在线率、最后采集时间和异常阈值来源。"],
        ["施工计划、值班、巡查", "按业务提交 / 日更新", "标注计划日期与数据来源，避免与实时数据混淆。"],
    ], [2500, 2200, 4800])
    add_callout(doc, "安全态势评分建议", "将事件严重度、交通影响、气象等级、结构物异常、设备故障和处置时效归一为 0–100 分；大屏展示等级与贡献项，不建议只展示单一分数。", RED)

    page_break(doc)
    h1(doc, "8. 开发验收清单")
    table(doc, ["验收维度", "验收点"], [
        ["画布与拼接", "以 9568 × 2808 无滚动展示；核心信息不压缝；在实际拼接屏上完成色彩与亮度校准。"],
        ["可读性", "3 米外可识别主标题、核心数字与红橙黄告警；非核心文字在运维距离可读。"],
        ["信息一致性", "同一事件在左栏、地图、右栏的编号、等级、位置和时间一致；筛选条件全局可见。"],
        ["状态完整性", "每个专题均具备默认、加载、空态、异常、权限不足五类状态，不出现空白区域。"],
        ["风险语义", "红/橙/黄/绿的含义统一；不能仅依靠颜色表达状态；色盲模式下仍可用图标或文字区分。"],
        ["性能", "首屏骨架快速呈现；地图操作不阻塞右栏；高频推送采用增量更新，避免整页闪烁。"],
        ["数据口径", "每项统计可查看时间范围、来源、刷新时间和口径说明；异常数据保留降级提示。"],
        ["操作闭环", "事件、管制、结构物、设备均可从总览定位至详情，并能进入视频或处置流程。"],
    ], [1900, 7660], "EAF5FC", 9.4)
    h2(doc, "8.1 交付物建议")
    add_bullets(doc, [
        "9568 × 2808 像素高保真主概览页及右侧 6 个专题状态稿。",
        "交通流及事件、养护巡查及运营两套全屏专题高保真稿。",
        "设计 Token、图表色板、图标库、组件状态与前端标注。",
        "地图图层字典、对象字段字典、事件等级与告警文案规范。",
    ])
    add_callout(doc, "本次梳理边界", "本文档基于提供的 G70 参考界面归纳为‘陕西交控路网安全态势感知’的 UI 方案；具体业务指标阈值、路网数据源、接口口径和实际屏体拼缝参数需在实施阶段与业务、数据及硬件团队联合确认。", AMBER)

    doc.core_properties.title = "陕西交控路网安全态势感知大屏 UI 设计文档"
    doc.core_properties.subject = "9568×2808 超宽大屏 UI 设计规范"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
