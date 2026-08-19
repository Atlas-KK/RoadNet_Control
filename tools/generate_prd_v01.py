from __future__ import annotations

import argparse
import os
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
LIGHT_YELLOW = "FFF2CC"
LIGHT_RED = "FCE4D6"
LIGHT_GREEN = "E2F0D9"
TEXT = "1F2937"
RED = "C00000"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    set_cell_width_dxa(cell, int(round(width_cm * 567)))


def set_cell_width_dxa(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_cm: list[float], indent_dxa: int = 90) -> list[int]:
    max_width_cm = 16.0
    total = sum(widths_cm)
    scale = min(1.0, max_width_cm / total) if total else 1.0
    widths = [width * scale for width in widths_cm]
    target_dxa = int(round(sum(widths) * 567))
    widths_dxa = [int(round(width * 567)) for width in widths]
    if widths_dxa:
        widths_dxa[-1] += target_dxa - sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(target_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "CBD5E1")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width_dxa))
        grid.append(col)
    return widths_dxa


def add_abstract_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "" if kind == "bullet" else "（%1）")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    return abstract_id


def new_numbering_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在打开 Word 后自动更新；如未更新，请右键选择“更新域”。"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2, placeholder, fld_char3])


def remove_template_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    styles = doc.styles
    existing_names = {style.name for style in styles}
    for style_name, style_type in (
        ("Normal", WD_STYLE_TYPE.PARAGRAPH),
        ("Subtitle", WD_STYLE_TYPE.PARAGRAPH),
        ("List Bullet", WD_STYLE_TYPE.PARAGRAPH),
        ("List Bullet 2", WD_STYLE_TYPE.PARAGRAPH),
        ("List Number", WD_STYLE_TYPE.PARAGRAPH),
        ("Table Grid", WD_STYLE_TYPE.TABLE),
    ):
        if style_name not in existing_names:
            styles.add_style(style_name, style_type)
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(3)

    heading_styles = {}
    for level, size, color in (
        (1, 18, BLUE),
        (2, 14, BLUE),
        (3, 11.5, "365F91"),
    ):
        style = next(style for style in styles if style.name == f"Heading {level}")
        heading_styles[level] = style
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
    doc._prd_heading_styles = heading_styles

    header = section.header.paragraphs[0]
    header.text = "路网智能管控｜需求规格说明书 V0.1"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string("7F8C8D")
    add_page_number(section.footer.paragraphs[0])
    doc._prd_bullet_abstract = add_abstract_numbering(doc, "bullet")
    doc._prd_decimal_abstract = add_abstract_numbering(doc, "decimal")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text)
    p.style = doc._prd_heading_styles[min(level, 3)]
    keep_with_next(p)


def add_para(doc: Document, text: str, bold_prefix: str | None = None, color: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def add_bullets(doc: Document, items: list[str], level: int = 0) -> None:
    num_id = new_numbering_instance(doc, doc._prd_bullet_abstract)
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, num_id)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    num_id = new_numbering_instance(doc, doc._prd_decimal_abstract)
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, num_id)
        p.add_run(item)


def add_callout(doc: Document, title: str, text: str, fill: str = LIGHT_YELLOW) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_dxa = set_table_geometry(table, [16.0], indent_dxa=150)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width_dxa(cell, widths_dxa[0])
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 110, 150, 110, 150)
    p = cell.paragraphs[0]
    p.add_run(f"{title}：").bold = True
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None,
              font_size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_dxa = set_table_geometry(table, widths) if widths else None
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        cell.text = value
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        if widths_dxa:
            set_cell_width_dxa(cell, widths_dxa[idx])
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(font_size)
    for r_idx, values in enumerate(rows):
        row = table.add_row()
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if widths_dxa:
                set_cell_width_dxa(cell, widths_dxa[c_idx])
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths_dxa = set_table_geometry(table, [16.0], indent_dxa=140)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width_dxa(cell, widths_dxa[0])
    set_cell_shading(cell, "F4F6F8")
    set_cell_margins(cell, 120, 140, 120, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("263238")


def add_fp(doc: Document, fp_id: str, name: str, goal: str, actors: str, pre: str, main: str,
           alternate: str, data: str, state: str, acceptance: str) -> None:
    add_heading(doc, f"{fp_id} {name}", 3)
    add_table(
        doc,
        ["要素", "规格"],
        [
            ["目标", goal],
            ["参与者", actors],
            ["前置条件", pre],
            ["主流程", main],
            ["异常/替代流程", alternate],
            ["数据读写", data],
            ["状态变化", state],
            ["验收要点", acceptance],
        ],
        [3.0, 13.0],
        8.5,
    )


def build_document(template: Path, output: Path) -> None:
    doc = Document(str(template))
    remove_template_body(doc)
    configure_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    r = p.add_run("路网智能管控")
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(32)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("需求规格说明书")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string("334155")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("版本 V0.1｜阶段一演示基线")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("64748B")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(160)
    p.add_run("编制日期：2026年8月18日\n适用终端：Windows 单机、本地浏览器、1920×1080、100% 缩放\n文档状态：评审稿").font.size = Pt(11)
    doc.add_page_break()

    add_heading(doc, "文档信息", 1)
    add_table(doc, ["项目", "内容"], [
        ["文档名称", "路网智能管控需求规格说明书"],
        ["版本", "V0.1（评审稿，未经用户定版不得升级为 V1.0）"],
        ["适用阶段", "阶段一：本地演示与功能验证"],
        ["产品负责人", "待指派"],
        ["技术负责人", "待指派"],
        ["业务负责人", "待指派"],
        ["数据安全负责人", "待指派"],
        ["编制依据", "现有项目功能、演示开发规格、五类案例、用户阶段确认结论"],
    ], [4.0, 12.0])
    add_heading(doc, "修订记录", 2)
    add_table(doc, ["版本", "日期", "修订说明", "修订人"], [
        ["V0.1", "2026-08-18", "形成17章需求规格说明书；纳入四类智能决策底座、设备指令联动、实时讲解、通行能力及三流审查。", "Codex（待项目组复核）"],
    ], [2.0, 3.0, 9.0, 2.0])
    add_callout(doc, "阶段边界", "阶段一全部业务数据、设备回执、人员回执与控制效果均为模拟数据。地图底图与可选大模型服务可作为展示支撑，但不得传入真实敏感业务数据；不接入真实业务系统和物理设备。")
    add_heading(doc, "目录", 1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    # 1
    add_heading(doc, "1 项目概述", 1)
    add_heading(doc, "1.1 建设背景", 2)
    add_para(doc, "高速公路突发事件处置涉及事件核实、交通影响研判、方案编制、资源协调、指令下发、设备联动和复盘归档等多环节。传统方式容易出现信息分散、处置依据不透明、措施协同困难和过程难追溯等问题。本项目建设“路网智能管控”演示系统，以事件为主线，将 GIS 数字孪生、交通流计算、智能策略生成、统一处置时序、设施监测和审计归档整合到同一工作台。")
    add_heading(doc, "1.2 产品定位", 2)
    add_para(doc, "产品采取分阶段建设。阶段一以可演示、可验证、可追溯为目标，验证从事件上报到处置闭环的关键业务链路；阶段二及以后再推进真实数据接入、多角色权限、生产安全、运维保障及实际设备控制。阶段一的所有控制均为模拟控制，不产生真实道路控制效果。")
    add_heading(doc, "1.3 核心目标", 2)
    add_table(doc, ["编号", "目标", "阶段一可验证结果"], [
        ["G-01", "事件全过程一屏协同", "完成上报、融合、研判、策略、确认、回执、监测、结案和审计闭环。"],
        ["G-02", "生成可比较、可解释、可追溯的策略", "候选策略同步展示适用条件、措施、资源、效果、风险、置信依据和推理来源。"],
        ["G-03", "统一确认与执行反馈", "高风险措施人工确认后模拟下发，并按系统、人员、设备展示回执。"],
        ["G-04", "现场态势动态可视", "GIS、实时讲解、通行能力、1km设施网格与设备状态同步变化。"],
        ["G-05", "形成可复盘资产", "支持审计记录、数据集导出和结案报告。"],
    ], [1.7, 5.0, 9.3])
    add_heading(doc, "1.4 非目标与后续范围", 2)
    add_bullets(doc, [
        "阶段一不接入真实交警、路政、消防、救援、收费、视频、信号或诱导发布系统。",
        "阶段一不向真实可变情报板、信号灯、车道指示灯、摄像头等物理设备下发控制。",
        "阶段一不建设生产级多租户、组织权限、单点登录、等保、安全运营与灾备体系。",
        "生产 SLA、RTO、RPO、并发量、容量和等保等级均待后续阶段确认。",
        "阶段一不以模型输出替代指挥员决策；涉及封道、分流、限速或物理设备控制的措施必须人工确认。",
    ])
    add_heading(doc, "1.5 术语与约定", 2)
    add_table(doc, ["术语", "定义"], [
        ["事件", "影响路网运行并需要研判或处置的业务对象。"],
        ["事件事实", "描述时间、地点、方向、占道、人员、车辆、环境等的结构化信息。"],
        ["候选策略", "由智能底座生成、可供比较选择的管控方案组合。"],
        ["处置措施", "策略中的可执行操作项，如全幅封道、分流诱导、视频巡查。"],
        ["统一处置时序", "按依赖与优先级组织措施确认、下发、反馈和升级的时序视图。"],
        ["事故点桩号", "事故所在道路里程标识，例如 G65 K1177.2。"],
        ["模拟数据", "为演示构造的事件、交通流、设备、人员和回执数据，不代表真实生产状态。"],
    ], [4.0, 12.0])

    # 2
    add_heading(doc, "2 用户与使用场景", 1)
    add_heading(doc, "2.1 用户角色", 2)
    add_table(doc, ["角色ID", "角色", "职责", "阶段一权限"], [
        ["U-01", "值班指挥员", "查看态势、发起或核实事件、比较策略、确认高风险措施、查看回执、结案与审计。", "系统内全部演示操作；不具备真实外部控制权限。"],
    ], [2.0, 3.0, 7.0, 4.0])
    add_callout(doc, "待确认", "产品负责人、技术负责人、业务负责人、数据安全负责人尚未指派；多角色权限矩阵在后续生产化阶段补充。", LIGHT_RED)
    add_heading(doc, "2.2 核心场景", 2)
    add_table(doc, ["场景", "触发", "用户目标", "完成标志"], [
        ["SC-01 事件上报与融合", "人工上报或加载案例", "快速形成结构化事实并识别重复事件", "事件进入研判队列。"],
        ["SC-02 智能研判", "事件事实可用", "理解影响、查看交通流计算与证据", "生成透明推理链和候选策略。"],
        ["SC-03 策略比较确认", "候选策略生成", "比较效果、风险、资源和适用条件", "选定方案版本并确认措施。"],
        ["SC-04 统一下发与联动", "措施被确认", "掌握指令是否到达系统、人员、设备", "显示回执并同步成功设备状态。"],
        ["SC-05 动态监测与修订", "事件处置中", "观察讲解、通行能力、队尾和设备", "必要时生成新版本或撤销失效措施。"],
        ["SC-06 结案复盘", "影响解除", "形成报告与可追溯资料", "事件关闭、报告可查看、审计可导出。"],
    ], [3.3, 3.5, 5.5, 4.0])
    add_heading(doc, "2.3 用户旅程", 2)
    add_table(doc, ["阶段", "用户行为", "系统反馈", "风险控制"], [
        ["发现", "加载案例或填写事件", "结构化提取、地图定位、融合提示", "原文保留；字段可复核。"],
        ["研判", "浏览推理与交通影响", "五步推演、图谱关系、规则与交通流结果", "证据来源与置信度可查看。"],
        ["决策", "比较并选择策略", "方案 V1/V2、预期效果、资源与风险", "高风险措施必须确认。"],
        ["执行", "逐项确认处置措施", "下发中、成功、部分成功或失败回执", "失败不改变设备原状态。"],
        ["监测", "查看态势和执行效果", "30秒讲解、通行能力、网格和地图同步", "异常提示；可修订方案。"],
        ["复盘", "结案与导出", "最终报告、JSONL 数据集、审计记录", "保留版本链和操作时间。"],
    ], [2.4, 4.4, 6.0, 4.0])

    # 3
    add_heading(doc, "3 需求总览", 1)
    requirements = [
        ["R-01", "P0", "事件接入与融合", "支持人工上报、案例加载、大模型结构化提取、事件合并与优先级排序。", "FP-01"],
        ["R-02", "P0", "GIS数字孪生", "在地图呈现事件、道路、队列、设备、资源、风险区和分流线路，并支持联动定位。", "FP-02"],
        ["R-03", "P0", "透明推理", "以五步推理展示事实、计算、知识、规则和结论之间的关系。", "FP-03"],
        ["R-04", "P0", "智能决策底座", "建设大模型、知识图谱、知识库、业务规则库四类能力，并可测试其输入、输出和依据。", "FP-04"],
        ["R-05", "P0", "候选策略比较", "生成多套方案并统一展示条件、措施、资源、效果、风险和置信依据。", "FP-05"],
        ["R-06", "P0", "方案版本与确认", "支持 V1/V2 版本、措施确认/驳回/失效/超时升级及依赖传播。", "FP-06"],
        ["R-07", "P0", "指令下发回执", "按系统、人员、设备模拟下发，展示状态、耗时、时间和失败原因。", "FP-07"],
        ["R-08", "P0", "设备状态联动", "成功或部分成功回执驱动网格和 GIS 设备状态同步，失败保持原状态。", "FP-08"],
        ["R-09", "P0", "基础设施网格", "按1km桩号连续分格，支持5/10/20km、筛选、事故高亮、横向拖拽和固定首列。", "FP-09"],
        ["R-10", "P0", "实时态势讲解", "每30秒生成可追溯态势简报，最多保留12条，并在模型失败时本地降级。", "FP-10"],
        ["R-11", "P0", "上下游通行能力", "展示正常/实时通行能力、保持率、拥堵等级和趋势，关联事件桩号。", "FP-11"],
        ["R-12", "P1", "事件孪生", "模拟车辆、雷达视频覆盖、轨迹和队尾演化。", "FP-12"],
        ["R-13", "P0", "审计与归档", "形成结案报告、操作审计和JSONL数据集，并关联事件与版本。", "FP-13"],
        ["R-14", "P0", "演示运行", "提供五类案例、本地启动、配置校验和自动化测试。", "FP-14"],
    ]
    add_table(doc, ["编号", "优先级", "需求名称", "需求说明", "功能映射"], requirements, [1.4, 1.3, 3.0, 8.5, 1.8], 8)
    add_heading(doc, "3.2 非功能需求", 2)
    add_table(doc, ["编号", "类别", "要求", "验收口径"], [
        ["NFR-01", "兼容", "Windows 单机、本地浏览器、1920×1080、浏览器100%缩放。", "主工作台无遮挡、无整页横向溢出。"],
        ["NFR-02", "布局", "设施网格横向超宽不得挤压上下游通行能力及右侧时序区。", "网格内部滚动；左侧设备类型列固定；相邻图表不变形。"],
        ["NFR-03", "可靠", "大模型简报返回HTTP 400、超时或结构不合法时启用本地规则降级。", "界面持续显示有效简报并标注来源。"],
        ["NFR-04", "可测试", "lint、单元测试、类型检查、生产构建均通过。", "基线为38个测试文件、182项测试；交付时执行并记录实际结果。"],
        ["NFR-05", "可追溯", "策略、措施、回执、设备变化与审计记录可关联。", "任一设备同步状态可反查具体措施与下发时间。"],
        ["NFR-06", "安全边界", "不得把模拟控制误标为真实控制；不得向真实设备下发。", "页面显示演示/模拟标识，接口清单无真实控制端点。"],
        ["NFR-07", "生产指标", "生产SLA、RTO、RPO、等保等级、并发和容量。", "待后续阶段确认。"],
    ], [1.6, 2.0, 8.4, 4.0], 8)

    # 4
    add_heading(doc, "4 信息架构与页面结构", 1)
    add_heading(doc, "4.1 页面结构", 2)
    add_table(doc, ["层级", "区域", "内容"], [
        ["全局", "顶部导航", "交控Logo、路网智能管控标题、系统在线、设置、数据集、审计。"],
        ["一级工作台", "事件列表", "活跃事件、演示案例加载、人工上报。"],
        ["一级工作台", "GIS数字孪生", "路网态势、事件孪生、现场视频、2D/全线视角及全屏。"],
        ["一级工作台", "态势与能力", "事件实时态势感知、上下游通行能力。"],
        ["一级工作台", "智能处置时序", "推理、预案、确认、措施状态、下发回执。"],
        ["一级工作台", "设施监测网格", "1km桩号网格、范围切换、设备筛选、设备卡片与地图联动。"],
        ["弹层", "报告/审计/数据集/设置", "结案报告、日志明细、JSONL导出、模型与演示配置。"],
    ], [2.5, 4.0, 10.0])
    add_heading(doc, "4.2 页面与需求映射", 2)
    add_table(doc, ["页面/区域", "核心需求", "关键状态"], [
        ["事件列表", "R-01、R-14", "空、提取中、待研判、处理中、已关闭"],
        ["GIS数字孪生", "R-02、R-08、R-12", "概览、事件聚焦、设备高亮、孪生演化"],
        ["实时态势感知", "R-03、R-10", "生成中、模型生成、本地降级、失败提示"],
        ["上下游通行能力", "R-11", "正常、下降、拥堵、采样中"],
        ["统一处置时序", "R-04～R-08", "待确认、已确认、下发中、成功、部分成功、失败、失效"],
        ["设施网格", "R-08、R-09", "5/10/20km、类型筛选、事故列、已同步、未同步"],
    ], [4.3, 7.0, 5.0])

    # 5 diagrams
    add_heading(doc, "5 业务流程与交互时序", 1)
    add_heading(doc, "5.1 事件闭环业务流程", 2)
    add_code_block(doc, """
flowchart LR
  A["事件上报/案例加载"] --> B["语义理解与结构化"]
  B --> C{"重复事件?"}
  C -- 是 --> D["事件融合/更新事实"]
  C -- 否 --> E["创建事件"]
  D --> F["五步透明推理"]
  E --> F
  F --> G["四类智能底座检索与校验"]
  G --> H["交通流计算"]
  H --> I["生成并比较候选策略"]
  I --> J{"措施风险级别"}
  J -- 低风险 --> K["自动执行模拟刷新/提示"]
  J -- 高风险 --> L["值班指挥员确认"]
  L --> M["模拟下发至系统/人员/设备"]
  K --> N["态势、能力、GIS与网格监测"]
  M --> N
  N --> O{"事实或效果变化?"}
  O -- 是 --> P["生成新方案版本/撤销失效措施"]
  P --> I
  O -- 否且影响解除 --> Q["结案报告/审计/数据集"]
""")
    add_heading(doc, "5.2 智能策略生成与执行时序", 2)
    add_code_block(doc, """
sequenceDiagram
  actor C as 值班指挥员
  participant UI as 工作台
  participant LLM as 大模型
  participant KG as 知识图谱
  participant KB as 知识库
  participant BR as 业务规则库
  participant TF as 交通流模型
  participant DS as 模拟下发服务
  C->>UI: 上报/选择事件
  UI->>LLM: 语义理解与结构化
  UI->>KG: 关联道路/设备/资源/风险/措施
  UI->>KB: 检索预案/案例/规范/经验
  UI->>BR: 校验适用条件/冲突/边界/安全
  UI->>TF: 计算瓶颈能力/密度/排队/队尾
  UI-->>C: 候选策略与证据比较
  C->>UI: 确认高风险措施
  UI->>DS: 模拟下发至系统/人员/设备
  DS-->>UI: 成功/部分成功/失败回执
  UI-->>C: 更新时序、GIS、网格、讲解和审计
""")
    add_heading(doc, "5.3 措施下发状态流", 2)
    add_code_block(doc, """
stateDiagram-v2
  [*] --> pending: 操作项生成
  pending --> dispatching: 指挥员确认或低风险自动触发
  pending --> rejected: 指挥员驳回
  pending --> obsolete: 上游事实变化/TMS撤销
  dispatching --> success: 全部目标成功
  dispatching --> partial_success: 部分目标成功
  dispatching --> failed: 全部目标失败/超时
  partial_success --> dispatching: 重新下发失败目标
  failed --> dispatching: 重试
  success --> obsolete: 事实变化导致措施失效
""")

    # 6
    add_heading(doc, "6 功能模块设计", 1)
    add_table(doc, ["模块ID", "模块", "职责", "主要输入", "主要输出"], [
        ["M-01", "事件中心", "上报、提取、融合、排序、聚焦与关闭事件。", "事件文本/案例", "Event、EventFact"],
        ["M-02", "GIS数字孪生", "统一呈现路网、事件、设施、资源、队列和联动状态。", "事件/设备/交通数据", "地图图层与高亮"],
        ["M-03", "智能决策底座", "大模型、知识图谱、知识库、规则库协同。", "事件事实/知识/规则", "推理证据与候选措施"],
        ["M-04", "交通流计算", "计算能力、密度、排队和队尾演化。", "车道/流量/速度/占用", "指标与趋势"],
        ["M-05", "策略中心", "生成、比较、选择和版本化策略。", "证据/计算/资源", "CandidateStrategy、Plan"],
        ["M-06", "统一处置时序", "组织措施依赖、确认、下发和状态。", "Plan/Measure", "Dispatch、Receipt"],
        ["M-07", "设施监测", "1km网格、筛选、状态显示与地图联动。", "Device/Effect", "网格与设备卡片"],
        ["M-08", "态势讲解", "周期生成态势摘要并在异常时本地降级。", "事件/交通/处置状态", "Narrative"],
        ["M-09", "能力监测", "显示上下游门架能力与变化趋势。", "TrafficSample", "能力卡片与曲线"],
        ["M-10", "审计与归档", "记录操作、导出数据集、生成结案报告。", "全流程数据", "Audit/Report/JSONL"],
    ], [1.7, 3.0, 5.0, 3.5, 3.5], 8)
    add_heading(doc, "6.2 功能点清单", 2)
    add_table(doc, ["功能点", "名称", "所属模块", "优先级", "说明"], [
        [f"FP-{i:02d}", name, module, priority, desc] for i, (name, module, priority, desc) in enumerate([
            ("事件上报与融合", "M-01", "P0", "文本提取、重复识别、合并和优先级。"),
            ("GIS路网态势", "M-02", "P0", "图层、聚焦、设备定位与高亮。"),
            ("五步透明推理", "M-03", "P0", "展示事实、知识、计算、规则和结论。"),
            ("四类智能底座", "M-03", "P0", "四类能力协同并可验收。"),
            ("候选策略比较", "M-05", "P0", "多方案同口径比较。"),
            ("方案版本与措施确认", "M-05/M-06", "P0", "版本、确认、驳回、失效与升级。"),
            ("指令下发与回执", "M-06", "P0", "系统、人员、设备的模拟回执。"),
            ("设备同步联动", "M-02/M-07", "P0", "成功设备更新网格与GIS。"),
            ("1km设施网格", "M-07", "P0", "范围、筛选、滚动与事故列。"),
            ("实时态势讲解", "M-08", "P0", "30秒刷新、历史与降级。"),
            ("上下游通行能力", "M-09", "P0", "能力、保持率、等级和趋势。"),
            ("事件孪生演化", "M-02", "P1", "车辆、覆盖、轨迹和队尾。"),
            ("审计与结案", "M-10", "P0", "报告、审计、JSONL。"),
            ("演示与质量校验", "M-01/M-10", "P0", "案例、配置、自动化测试。"),
        ], start=1)
    ], [1.7, 4.0, 3.0, 1.5, 6.0], 8)

    # 7 detailed functions
    add_heading(doc, "7 详细功能规格", 1)
    add_heading(doc, "7.1 功能点七要素规格", 2)
    add_fp(doc, "FP-01", "事件上报与融合", "将非结构化事件转为可研判对象并防止重复处置。", "值班指挥员、事件中心、大模型",
           "系统已加载；用户可选择案例或录入文本。",
           "1) 提交事件；2) 语义提取道路、桩号、方向、占道、车辆、人员、环境；3) 计算相似度；4) 新建或提示合并；5) 更新排序与GIS焦点。",
           "模型不可用或结构不合法时使用本地规则解析；缺失字段标注待核实；用户可保留为独立事件。",
           "读：案例、既有事件；写：Event、EventFact、AuditLog。",
           "draft→analyzing→active；或 incoming→merged。",
           "原文、结构化结果和合并依据可查看；重复事件不会无提示覆盖。")
    add_fp(doc, "FP-02", "GIS路网态势", "在统一空间视图呈现与事件有关的道路、设备、资源和风险。", "值班指挥员、GIS模块",
           "事件或全网模拟数据已加载。",
           "展示路网、事件点、队列、摄像头、VMS、车道指示灯、信号灯、资源、雾区、隧道与分流线路；点击对象显示详情；设备卡片可反向定位。",
           "地图服务不可用时保留业务面板并提示底图不可用；不得阻断策略与回执演示。",
           "读：Event、Device、Resource、TrafficSample；写：MapFocus、AuditLog。",
           "overview↔event_focus↔device_focus。",
           "网格卡片点击后地图定位并高亮；设备状态与网格一致。")
    add_fp(doc, "FP-03", "五步透明推理", "让指挥员理解策略如何由事实和依据产生。", "值班指挥员、智能决策底座",
           "事件事实已结构化。",
           "按事件理解、信息汇聚、依据匹配、影响推演、策略形成五步展示；每步列出输入、过程、证据和结论。",
           "证据不足时明确标注“不确定/待核实”，不得伪造确定性依据。",
           "读：EventFact、KnowledgeItem、GraphRelation、BusinessRule、TrafficCalc；写：ReasoningTrace。",
           "pending→generated；事实更新后 previous→superseded。",
           "可从策略追溯至规则、知识条目、交通计算和事件事实。")
    add_fp(doc, "FP-04", "四类智能决策底座", "通过大模型、知识图谱、知识库、业务规则库形成可比较、可解释、可追溯的策略。", "智能决策底座",
           "模拟事件和阶段一知识资产可用。",
           "大模型完成语义理解、信息归纳和策略表达；知识图谱关联事件、道路、设施、资源、风险和措施；知识库检索预案、历史案例、行业规范与经验；规则库校验适用条件、资源冲突、执行边界与安全约束。",
           "任一底座不可用时标注降级来源；规则校验失败的措施不得进入可确认状态；知识来源缺失时降低置信说明。",
           "读：四类底座资产、EventFact；写：EvidenceRef、RuleCheck、ReasoningTrace。",
           "available/degraded/unavailable；规则结果 pass/warn/block。",
           "四类底座分别具有可演示输入输出；候选策略至少显示一项可追溯依据和规则校验结论。")
    add_fp(doc, "FP-05", "候选策略比较", "以统一口径提供多套管控策略，辅助选择而非替代决策。", "值班指挥员、策略中心",
           "透明推理、交通流计算、资源与规则校验已完成。",
           "生成多套候选策略；展示适用条件、实施措施、资源需求、预期效果、风险、置信依据；支持切换与选定。",
           "无可行策略时说明阻断规则和缺失条件；不自动确认高风险措施。",
           "读：ReasoningTrace、TrafficCalc、Resource、RuleCheck；写：CandidateStrategy、PlanVersion。",
           "generated→compared→selected；未选方案 retained/rejected。",
           "同一事件可比较至少两套候选策略；字段口径一致且理由可追溯。")
    add_fp(doc, "FP-06", "方案版本与措施确认", "管理事实变化下的版本、依赖和人工确认边界。", "值班指挥员、策略中心、时序引擎",
           "存在已选方案与措施依赖。",
           "展示方案V1/V2及差异；低风险态势刷新/信息提示可自动模拟执行；封道、分流、限速和物理设备控制必须逐项确认；支持驳回、失效、误报、超时升级。",
           "事实变化触发新版本并通过TMS传播撤销依赖措施；旧版本保留只读记录。",
           "读：PlanVersion、Measure、Dependency；写：Confirmation、MeasureState、AuditLog。",
           "draft→active→superseded/closed；measure pending→confirmed/rejected/obsolete。",
           "未确认高风险措施不得进入下发中；版本链和撤销原因可查询。")
    add_fp(doc, "FP-07", "指令下发与回执", "显示操作项是否正常下发至系统、人员和设备。", "值班指挥员、模拟下发服务",
           "措施已确认，或属于允许自动执行的低风险项。",
           "生成系统/人员/设备目标；显示待下发、下发中、成功、部分成功、失败；已成功显示“已下发·耗时”；展开查看目标类型、名称、状态、确认时间或失败原因。",
           "部分成功显示失败目标数量；失败目标允许模拟重试；超时按失败处理并保留原因。",
           "读：Measure、DispatchTarget；写：Dispatch、Receipt、AuditLog。",
           "pending→dispatching→success/partial_success/failed。",
           "三类目标均可演示；摘要、耗时、失败数和回执明细一致。")
    add_fp(doc, "FP-08", "设备状态同步联动", "让指令成功结果同步反映到网格和GIS。", "设施监测、GIS、模拟下发服务",
           "设备目标返回成功或部分成功。",
           "按成功目标写入设备效果；网格卡片显示受指令影响、指令名称、下发时间、当前状态；GIS标记同步更新。全幅封道使VMS显示红色“禁止通行”；分流诱导显示黄色“减速慢行·按指引分流”；车道管控显示红叉/绿箭头；信号控制显示黄闪管控；视频巡查显示调阅中/正常/异常。",
           "失败设备保持原状态并显示未同步/失败；人员或系统回执不无故改变设备状态。",
           "读：Receipt、Device；写：DeviceEffect、DeviceState、MapMarkerState。",
           "unchanged→syncing→synced；或 unchanged→sync_failed。",
           "成功目标同步、失败目标不变；网格和GIS状态、时间、关联措施一致。")
    add_fp(doc, "FP-09", "1km基础设施网格", "按道路桩号清晰展示事故上下游设施实际分布与状态。", "值班指挥员、设施监测",
           "聚焦事件具有道路和事故桩号。",
           "按整公里生成连续1km列；事故点所属列高亮并显示桩号；切换5/10/20km时展示上下游各对应范围；设备按桩号归格；筛选全部/VMS/信号灯/车道指示灯/摄像头/风机/风向传感器；内部横向拖拽和底部滚动条浏览，设备类型首列固定。",
           "筛选只隐藏设备卡片，不改变桩号列；超宽内容不得扩张外层布局；边界桩号无设备时仍显示空网格。",
           "读：Event、Device、DeviceState；写：GridRange、DeviceFilter、MapFocus。",
           "range 5/10/20；filter all/type；accident-column persistent。",
           "20km完整加载且图表不变形；事故列清晰；设备桩号归属正确。")
    add_fp(doc, "FP-10", "实时态势讲解", "周期归纳事件变化、交通影响和处置进度。", "值班指挥员、大模型、本地规则引擎",
           "聚焦事件存在。",
           "每30秒基于事实、交通和处置状态生成讲解；展示生成时间和来源；历史最多12条；内容覆盖事件概况、交通影响、处置进展和关注事项。",
           "模型HTTP 400、超时、空响应或Schema不合法时，使用本地规则生成；首次无数据时显示明确引导。",
           "读：Event、TrafficSample、Measure、Receipt；写：Narrative、AuditLog。",
           "idle→generating→model_success/local_fallback；history capped at 12。",
           "模型失败不出现空白区；降级简报可读、来源明确、刷新周期正确。")
    add_fp(doc, "FP-11", "上下游通行能力", "量化事件对上下游门架能力的影响并展示趋势。", "值班指挥员、交通流计算",
           "聚焦事件存在并配置上下游模拟门架。",
           "展示门架桩号、正常能力、实时能力、保持率、下降比例、拥堵等级和历史趋势；随模拟采样更新。",
           "样本缺失时显示等待采样；异常值不进入趋势并记录告警。",
           "读：Gantry、TrafficSample；写：CapacityResult。",
           "sampling→normal/degraded/congested。",
           "上下游口径一致；保持率与能力值计算一致；布局不受设施网格影响。")
    add_fp(doc, "FP-12", "事件孪生演化", "演示事件周边车辆、感知覆盖、轨迹与队尾变化。", "值班指挥员、GIS孪生模块",
           "选择事件孪生视角。",
           "显示模拟车辆、雷达视频覆盖、轨迹、占道区和队尾；按时间推进位置和队列长度。",
           "性能不足时降低动画频率，保留关键事件和队尾状态。",
           "读：TwinScenario、TrafficSample；写：TwinFrame。",
           "stopped↔playing↔paused。",
           "演化可暂停/继续；队尾数据与态势讲解口径一致。")
    add_fp(doc, "FP-13", "审计与结案", "将处置全过程沉淀为可复盘、可导出的记录。", "值班指挥员、审计模块",
           "事件已形成处置过程。",
           "记录上报、事实更新、策略选择、措施确认、下发回执、设备同步、版本替换和关闭；生成结案报告；导出JSONL数据集。",
           "未关闭事件可生成阶段报告但不得标为最终结案；导出失败保留页面记录。",
           "读：全流程对象；写：AuditLog、FinalReport、DatasetExport。",
           "event active→closing→closed；report draft→final。",
           "报告可追溯至事件和方案版本；审计时间顺序完整。")
    add_fp(doc, "FP-14", "演示与质量校验", "保证阶段一可重复启动、演示和验证。", "开发/测试人员、值班指挥员",
           "Windows环境与项目依赖已准备。",
           "提供五类案例；本地启动后按脚本演示完整链路；执行lint、单元测试、类型检查和生产构建。",
           "模型或地图配置缺失时采用可识别降级，不影响核心闭环。",
           "读：配置、案例、测试；写：运行日志、测试报告。",
           "not_started→running→passed/failed。",
           "基线38个测试文件、182项测试；交付时记录实际命令和结果。")

    add_heading(doc, "7.2 核心状态模型", 2)
    add_table(doc, ["对象", "状态", "允许流转", "约束"], [
        ["事件", "draft/analyzing/active/closing/closed/merged", "draft→analyzing→active→closing→closed；incoming→merged", "closed只读；merged必须指向主事件。"],
        ["方案", "draft/active/superseded/closed", "draft→active→superseded；active→closed", "同一事件仅一个当前active版本。"],
        ["措施", "pending/confirmed/rejected/obsolete", "pending→confirmed/rejected/obsolete；confirmed→obsolete", "高风险未confirmed不得下发。"],
        ["下发", "pending/dispatching/success/partial_success/failed", "按5.3状态图流转", "终态需回执时间；部分成功需失败数。"],
        ["设备同步", "unchanged/syncing/synced/sync_failed", "unchanged→syncing→synced/sync_failed", "仅成功设备可写入新状态。"],
        ["简报", "idle/generating/model_success/local_fallback", "idle→generating→成功或降级", "最多保留12条。"],
    ], [2.4, 4.4, 5.2, 4.4], 8)
    add_heading(doc, "7.3 核心业务规则", 2)
    add_table(doc, ["规则ID", "规则", "处理"], [
        ["BR-01", "低风险态势刷新和信息提示可自动模拟执行。", "直接进入模拟执行并记录来源。"],
        ["BR-02", "封道、分流、限速和物理设备控制必须人工确认。", "未确认时保持pending，不触发设备变化。"],
        ["BR-03", "只有success或partial_success中的成功设备可同步。", "失败设备保持原状态并记录原因。"],
        ["BR-04", "VMS“禁止通行”显示红色，“减速慢行”显示黄色，未下发显示灰色。", "网格与GIS应用同一颜色映射。"],
        ["BR-05", "事故设施网格按整公里对齐，每列代表连续1km。", "事故点落入floor(桩号)～floor(桩号)+1列。"],
        ["BR-06", "5/10/20km表示事故点上下游各对应距离。", "切换后更新列和设备，不改变事故点归属。"],
        ["BR-07", "模型失败时不得让实时态势区域空白。", "转本地规则并标注local_fallback。"],
        ["BR-08", "事实变化导致措施不再适用时，通过依赖关系传播obsolete。", "生成新方案版本，旧版保留。"],
        ["BR-09", "所有数据、回执和控制效果均标识为模拟。", "不得用生产术语暗示真实执行。"],
    ], [2.0, 10.0, 4.5], 8)

    # 8
    add_heading(doc, "8 界面与交互设计", 1)
    add_heading(doc, "8.1 总体布局", 2)
    add_para(doc, "主工作台采用事件列表、GIS与设施区、态势与能力区、智能处置时序区的多区域布局。1920×1080、100%缩放为阶段一验收基准。各区域必须通过容器宽度约束独立滚动，任何内部宽表不得扩展页面网格轨道或挤压相邻模块。")
    add_table(doc, ["区域", "交互要求", "防溢出要求"], [
        ["设施监测网格", "5/10/20km切换、类型筛选、卡片点击、横向拖拽、底部滚动。", "外层min-width:0；滚动容器overflow-x:auto；固定首列；桩号列定宽。"],
        ["上下游通行能力", "曲线/卡片同步采样、折叠或全屏。", "自身容器宽度100%；图表resize；不得由网格内容撑宽。"],
        ["统一处置时序", "滚动查看操作项、展开回执、确认高风险措施。", "仅纵向内部滚动；长目标名省略并可查看完整值。"],
        ["GIS", "图层切换、2D/全线视角、全屏、设备聚焦。", "地图实例随容器resize；不得覆盖下方网格。"],
    ], [3.3, 7.0, 6.0])
    add_heading(doc, "8.2 状态与颜色语义", 2)
    add_table(doc, ["对象/状态", "颜色语义", "显示要求"], [
        ["事故点网格", "事故红/浅红背景", "表头与整列高亮，不遮挡设备。"],
        ["VMS 禁止通行", "红色", "文本与真实语义一致。"],
        ["VMS 减速慢行", "黄色", "保证浅色背景对比度。"],
        ["VMS 未下发", "灰色", "显示“未下发内容”。"],
        ["下发成功", "绿色", "显示成功和耗时。"],
        ["部分成功", "橙色", "显示失败目标数量。"],
        ["下发失败/未同步", "红色", "显示失败原因入口。"],
        ["受指令影响", "蓝色边框/已同步标签", "显示关联指令和更新时间。"],
    ], [4.5, 3.5, 8.3])
    add_heading(doc, "8.3 空态、加载与异常", 2)
    add_bullets(doc, [
        "未选事件：实时讲解、通行能力、设施网格和处置时序显示引导文案，不显示伪造结果。",
        "生成中：显示明确加载状态和最近一次有效结果，避免内容闪空。",
        "模型异常：提示模型异常原因的简化信息，同时使用本地规则简报。",
        "地图异常：保留非地图业务功能，提示检查地图配置。",
        "无设备：保留1km桩号列并显示空状态，不折叠桩号结构。",
    ])

    # 9
    add_heading(doc, "9 数据设计", 1)
    add_heading(doc, "9.1 数据边界", 2)
    add_para(doc, "阶段一数据全部为模拟数据，逻辑数据模型用于前端状态和演示流程，不承诺生产数据库结构。外部地图底图与可选大模型服务只用于展示支撑；传给模型的数据必须是虚构或脱敏演示数据。")
    add_heading(doc, "9.2 核心数据对象", 2)
    add_table(doc, ["对象", "关键字段", "来源", "主要消费者"], [
        ["Event", "id, route, stake, direction, severity, status, createdAt", "上报/案例", "GIS、策略、讲解、审计"],
        ["EventFact", "eventId, category, value, confidence, source, verified", "模型/本地解析", "推理、图谱、规则"],
        ["KnowledgeItem", "id, type, title, source, version, content", "阶段一知识资产", "知识库检索、证据"],
        ["GraphRelation", "sourceId, relation, targetId, evidence", "知识图谱", "推理与依赖传播"],
        ["BusinessRule", "id, condition, action, severity, version", "规则库", "策略校验、状态约束"],
        ["TrafficCalc", "eventId, capacity, density, queueDensity, tail, assumptions", "交通流模型", "策略、能力、讲解"],
        ["CandidateStrategy", "id, conditions, measures, resources, effects, risks, confidenceRefs", "策略引擎", "策略比较"],
        ["PlanVersion", "id, eventId, version, status, selectedStrategyId, supersedes", "策略中心", "时序、审计"],
        ["Measure", "id, planId, type, riskLevel, dependencyIds, status", "策略中心", "确认与下发"],
        ["Dispatch", "id, measureId, status, startedAt, completedAt, elapsed", "模拟下发", "时序与审计"],
        ["DispatchTarget", "type, targetId, name, status, receiptAt, failureReason", "模拟下发", "回执明细"],
        ["Device", "id, type, route, stake, baseState, mapPosition", "模拟设施数据", "网格与GIS"],
        ["DeviceEffect", "deviceId, measureId, syncStatus, content, color, updatedAt", "成功回执", "网格、GIS、审计"],
        ["TrafficSample", "gantryId, timestamp, normalCapacity, actualCapacity, speed, flow", "模拟采样", "能力趋势与讲解"],
        ["Narrative", "eventId, generatedAt, source, sections, fallbackReason", "模型/本地规则", "实时态势区"],
        ["AuditLog", "timestamp, actor, action, objectType, objectId, before, after", "全模块", "审计与报告"],
    ], [3.2, 7.6, 3.0, 3.0], 7.7)
    add_heading(doc, "9.3 关键字段约束", 2)
    add_table(doc, ["字段", "类型/枚举", "约束"], [
        ["event.stake", "number", "公里值；用于事故网格归属；必须与route组合解释。"],
        ["dispatch.status", "pending/dispatching/success/partial_success/failed", "部分成功至少有1个成功和1个失败目标。"],
        ["target.type", "system/person/device", "阶段一至少可演示三类。"],
        ["device.type", "vms/signal/lane_indicator/camera/fan/wind_sensor", "筛选枚举与显示标签一一对应。"],
        ["effect.color", "red/yellow/gray/green/other", "VMS颜色必须符合BR-04。"],
        ["narrative.source", "model/local_fallback", "降级时fallbackReason必填。"],
        ["plan.version", "V1/V2/...", "同一事件递增且不可覆盖旧版本。"],
    ], [4.2, 5.3, 7.3])
    add_heading(doc, "9.4 数据生命周期", 2)
    add_table(doc, ["阶段", "新增/更新数据", "保留要求"], [
        ["事件接入", "Event、EventFact、AuditLog", "原始输入与结构化事实同时保留。"],
        ["策略生成", "ReasoningTrace、EvidenceRef、TrafficCalc、CandidateStrategy", "保留来源、假设与规则版本。"],
        ["人工确认", "PlanVersion、MeasureState、Confirmation", "保留操作者和时间。"],
        ["模拟下发", "Dispatch、DispatchTarget、Receipt、DeviceEffect", "失败原因与成功设备清单不可丢失。"],
        ["动态监测", "TrafficSample、Narrative、TwinFrame", "简报内存历史最多12条；归档口径待后续确认。"],
        ["结案", "FinalReport、DatasetExport、AuditLog", "阶段一保留周期待确认。"],
    ], [3.0, 8.0, 5.0])

    # 10
    add_heading(doc, "10 接口与集成", 1)
    add_heading(doc, "10.1 阶段一接口边界", 2)
    add_callout(doc, "明确约束", "阶段一不接入真实业务系统、真实人员通信系统或真实物理设备。以下接口均为前端内部逻辑接口或模拟服务。地图底图和可选模型API不属于业务控制接口。")
    add_table(doc, ["接口ID", "名称", "方向", "数据", "失败处理", "阶段"], [
        ["IF-01", "事件案例/人工上报", "UI→事件中心", "事件文本/案例ID", "校验并提示；保留输入", "阶段一模拟"],
        ["IF-02", "智能底座编排", "事件中心→智能底座", "事件事实/上下文", "底座状态降级并标注", "阶段一逻辑"],
        ["IF-03", "交通流计算", "策略中心→交通模型", "车道/流量/速度/占用", "拒绝异常样本，显示假设", "阶段一逻辑"],
        ["IF-04", "模拟下发", "时序引擎→模拟目标", "措施/目标清单", "返回部分成功/失败原因", "阶段一模拟"],
        ["IF-05", "设备效果同步", "回执→设施/GIS", "成功设备效果", "失败保持原状态", "阶段一逻辑"],
        ["IF-06", "实时讲解", "态势区→模型/本地规则", "演示事件上下文", "HTTP400/超时/Schema错则降级", "阶段一可选模型"],
        ["IF-07", "地图底图", "GIS→地图服务", "底图瓦片/脚本", "提示不可用，业务面板继续", "阶段一展示"],
        ["IF-08", "JSONL导出", "审计→本地下载", "过程样本", "提示失败并保留记录", "阶段一逻辑"],
    ], [1.6, 3.1, 3.0, 4.1, 4.2, 2.2], 7.5)
    add_heading(doc, "10.2 模型接口约束", 2)
    add_bullets(doc, [
        "输入只允许虚构或脱敏演示数据，不发送真实人员身份、真实设备凭证和生产控制参数。",
        "输出必须通过结构化Schema验证；不合格结果不得直接覆盖界面状态。",
        "模型输出仅用于语义理解、归纳和策略表达，业务规则与状态机拥有最终校验权。",
        "密钥不得写入源码或导出数据；配置方式和轮换策略待生产阶段完善。",
    ])
    add_heading(doc, "10.3 后续真实集成（待确认）", 2)
    add_table(doc, ["类别", "待确认内容"], [
        ["业务系统", "事件处置、诱导发布、视频监控、交通信号、收费、救援等系统清单、协议和责任边界。"],
        ["人员协同", "交警、路政、消防、养护、救援的组织编码、通知渠道和回执标准。"],
        ["设备控制", "设备协议、权限、双人复核、回滚、离线策略和安全隔离。"],
        ["数据治理", "数据分级分类、主数据、质量、留存、脱敏、跨域交换和审计。"],
    ], [3.4, 12.6])

    # 11
    add_heading(doc, "11 权限与安全", 1)
    add_heading(doc, "11.1 阶段一权限", 2)
    add_table(doc, ["功能", "值班指挥员", "约束"], [
        ["查看事件/GIS/态势/通行能力/网格", "允许", "全部为演示数据。"],
        ["人工上报/加载案例", "允许", "写入本地模拟状态。"],
        ["选择策略/确认措施", "允许", "只触发模拟下发。"],
        ["查看回执/重试失败目标", "允许", "不触达真实系统、人员或设备。"],
        ["结案/导出审计与数据集", "允许", "导出内容必须标注模拟。"],
        ["系统管理/用户管理/生产配置", "不在阶段一", "后续待设计。"],
    ], [7.0, 3.0, 6.0])
    add_heading(doc, "11.2 安全要求", 2)
    add_bullets(doc, [
        "界面、报告、导出数据明确标注“模拟数据/演示控制”。",
        "禁止在源代码、测试快照、文档和数据集中保存真实密钥。",
        "模型输出需经Schema、规则库和状态机三层约束后才能影响业务状态。",
        "高风险措施必须保留确认人、确认时间、关联版本和目标回执。",
        "生产等保等级、身份认证、最小权限、日志留存与灾备要求待后续阶段确认。",
    ])

    # 12
    add_heading(doc, "12 埋点、指标与审计", 1)
    add_table(doc, ["事件名", "触发时机", "核心属性", "用途"], [
        ["event_received", "事件提交/案例加载", "eventId, source, route, stake", "事件接入统计"],
        ["event_merged", "事件合并", "sourceId, targetId, reason", "融合效果复核"],
        ["strategy_generated", "候选策略完成", "eventId, count, evidenceCount, blockedCount", "策略质量"],
        ["strategy_selected", "用户选定方案", "eventId, strategyId, planVersion", "决策路径"],
        ["measure_confirmed", "确认高风险措施", "measureId, actor, timestamp", "责任追溯"],
        ["dispatch_result", "模拟下发结束", "status, elapsed, targetCounts", "执行反馈"],
        ["device_synced", "设备状态更新", "deviceId, measureId, syncStatus", "联动一致性"],
        ["narrative_generated", "简报生成", "source, latency, fallbackReason", "模型可用性"],
        ["plan_superseded", "新版本替代旧版", "oldVersion, newVersion, reason", "版本追溯"],
        ["event_closed", "事件结案", "eventId, duration, finalPlan", "闭环统计"],
    ], [3.8, 4.0, 6.0, 3.0], 8)
    add_heading(doc, "12.2 阶段一产品指标", 2)
    add_table(doc, ["指标", "口径", "阶段一目标"], [
        ["闭环可演示率", "五类案例中能完成上报→研判→确认→回执→监测→结案的案例数/5", "5/5"],
        ["高风险误自动下发数", "未人工确认即进入dispatching的高风险措施数", "0"],
        ["设备错误同步数", "失败设备状态被改变或网格/GIS不一致的数量", "0"],
        ["简报降级可用率", "模拟模型异常时仍产出有效本地简报的次数/异常次数", "100%"],
        ["自动化质量门禁", "lint、测试、类型、构建通过项", "4/4"],
        ["生产运营指标", "真实处置时长、策略采纳率、真实控制成功率等", "待真实数据接入后确认"],
    ], [4.2, 8.3, 4.3])

    # 13
    add_heading(doc, "13 非功能与运行保障", 1)
    add_table(doc, ["类别", "阶段一要求", "生产化待确认"], [
        ["部署", "Windows单机、本地浏览器运行。", "服务器拓扑、容器、网络区划。"],
        ["显示", "1920×1080、100%缩放；主界面无整页横向溢出。", "多分辨率、大屏拼接、移动端。"],
        ["刷新", "态势讲解每30秒；通行能力按模拟采样更新。", "真实采样频率和时钟同步。"],
        ["性能", "20km设施网格可完整加载并在内部横向浏览。", "并发量、响应时间、容量压测指标。"],
        ["可靠性", "模型异常本地降级；地图异常不阻断处置链。", "SLA、RTO、RPO、双活与容灾。"],
        ["安全", "模拟标识、密钥不入库、模型Schema与规则校验。", "等保等级、密码体系、安全运营。"],
        ["可维护", "TypeScript类型、模块化状态、自动化测试。", "监控告警、变更、灰度、回滚。"],
    ], [3.0, 7.0, 6.0])
    add_heading(doc, "13.2 故障降级矩阵", 2)
    add_table(doc, ["故障", "用户表现", "降级策略", "审计"], [
        ["大模型HTTP 400/超时", "提示已启用本地讲解", "本地规则生成，保留最近有效内容", "记录错误类型与fallback"],
        ["模型Schema不合法", "不展示不完整内容", "校验失败后本地生成", "记录校验字段"],
        ["地图底图失败", "GIS提示底图异常", "保留业务面板和非底图信息", "记录地图错误"],
        ["模拟下发失败", "操作项红色告警", "设备不变，可查看原因/重试", "记录目标级回执"],
        ["交通样本异常", "显示等待或最近有效值", "剔除异常样本，不误算趋势", "记录样本原因"],
    ], [4.0, 4.5, 5.2, 3.0])

    # 14
    add_heading(doc, "14 测试与验收", 1)
    add_heading(doc, "14.1 测试策略", 2)
    add_para(doc, "测试覆盖领域规则、状态流转、组件交互、布局边界、降级逻辑与生产构建。交付时应执行项目现有lint、单元测试、类型检查和生产构建，并在本章记录实际结果。当前基线为38个测试文件、182项测试。")
    add_heading(doc, "14.2 验收用例", 2)
    tests = [
        ["TC-01", "加载任一案例", "事件结构化、定位并进入研判", "R-01"],
        ["TC-02", "输入相似事件", "提示合并依据且不静默覆盖", "R-01"],
        ["TC-03", "生成策略", "可查看五步推理与四类底座依据", "R-03/R-04"],
        ["TC-04", "比较候选方案", "条件、措施、资源、效果、风险、置信依据齐全", "R-05"],
        ["TC-05", "未确认全幅封道", "不得进入下发，不改变设备", "R-06/R-08"],
        ["TC-06", "确认含三类目标的措施", "显示下发中及系统/人员/设备回执", "R-07"],
        ["TC-07", "模拟部分成功", "显示失败数；成功设备同步，失败设备不变", "R-07/R-08"],
        ["TC-08", "下发全幅封道", "VMS显示红色“禁止通行”，GIS同步", "R-08"],
        ["TC-09", "下发分流诱导", "VMS显示黄色“减速慢行·按指引分流”", "R-08"],
        ["TC-10", "切换5/10/20km", "上下游各对应范围的连续1km网格", "R-09"],
        ["TC-11", "在20km网格拖动", "内部可横向浏览，首列固定，通行能力图不变形", "R-09/NFR-02"],
        ["TC-12", "筛选设备类型", "仅卡片变化，桩号网格结构不变", "R-09"],
        ["TC-13", "模拟模型HTTP400", "显示本地降级简报，无空白", "R-10"],
        ["TC-14", "连续刷新简报", "30秒更新，最多12条，来源可见", "R-10"],
        ["TC-15", "查看通行能力", "正常/实时能力、保持率、等级、趋势一致", "R-11"],
        ["TC-16", "事实变化触发V2", "旧方案失效可追溯，依赖措施撤销", "R-06"],
        ["TC-17", "结案与导出", "报告、审计、JSONL关联同一事件与版本", "R-13"],
        ["TC-18", "运行质量命令", "lint、182项测试、类型检查、构建全部通过或记录最新实际值", "R-14"],
    ]
    add_table(doc, ["用例", "操作/条件", "预期结果", "映射"], tests, [1.5, 5.0, 8.0, 2.4], 8)
    add_heading(doc, "14.3 验收前置与证据", 2)
    add_bullets(doc, [
        "使用Windows单机、本地浏览器、1920×1080、100%缩放。",
        "使用五类演示案例及可复现的模拟失败数据。",
        "保留测试命令输出、关键页面截图、回执明细、审计导出和最终构建产物。",
        "任何待确认项不作为V0.1阶段一验收阻断项，但进入生产化前必须决策。",
    ])

    # 15
    add_heading(doc, "15 发布、部署与培训", 1)
    add_heading(doc, "15.1 阶段一发布流程", 2)
    add_numbered(doc, [
        "核对Node.js及项目依赖，按README完成本地配置。",
        "执行lint、单元测试、类型检查和生产构建，保存结果。",
        "启动本地服务，在1920×1080、100%缩放下完成关键页面检查。",
        "依次演示五类案例，核对策略、回执、设备联动、简报和通行能力。",
        "确认报告和审计导出可用，标记版本为V0.1评审稿。",
    ])
    add_heading(doc, "15.2 培训内容", 2)
    add_table(doc, ["对象", "内容", "产物"], [
        ["值班指挥员", "事件加载、策略比较、高风险确认、回执查看、设备联动、结案。", "演示手册/操作说明"],
        ["测试人员", "案例数据、失败注入、状态流转、布局、质量命令。", "测试用例与记录"],
        ["开发维护人员", "模块结构、模拟数据、智能底座、降级、状态模型。", "README/代码修改日志"],
    ], [3.0, 9.0, 4.0])
    add_heading(doc, "15.3 后续生产化门禁", 2)
    add_bullets(doc, [
        "完成四类负责人指派与生产责任边界确认。",
        "确认真实接口清单、数据分类分级、设备控制安全方案。",
        "确认多角色权限、双人复核、回滚和应急人工接管机制。",
        "确认SLA、RTO、RPO、并发、容量、等保与运维指标。",
        "完成真实环境集成、渗透测试、性能压测、容灾演练和上线审批。",
    ])

    # 16
    add_heading(doc, "16 风险与待确认事项", 1)
    add_table(doc, ["编号", "风险/待确认", "影响", "阶段一措施", "责任人"], [
        ["RK-01", "模型输出不稳定或HTTP错误", "简报/策略表达失败", "Schema校验+本地规则降级", "技术负责人（待指派）"],
        ["RK-02", "模拟状态被误解为真实执行", "业务误判", "全局模拟标识、报告声明、无真实控制接口", "产品/业务负责人（待指派）"],
        ["RK-03", "设施网格宽度挤压其他图表", "大屏变形", "内部滚动、固定首列、容器min-width:0与回归测试", "技术负责人（待指派）"],
        ["RK-04", "知识/规则版本不清", "依据不可追溯", "记录来源、版本、规则ID和证据引用", "业务负责人（待指派）"],
        ["RK-05", "生产控制缺少安全边界", "误控制风险", "阶段一禁止真实控制；生产化专项设计", "数据安全负责人（待指派）"],
        ["TBD-01", "产品/技术/业务/数据安全负责人", "评审与责任闭环", "待指派", "项目组"],
        ["TBD-02", "生产SLA/RTO/RPO/并发/容量", "无法形成生产验收基线", "后续阶段确认", "待指派"],
        ["TBD-03", "等保等级及安全合规", "无法完成生产安全设计", "后续阶段确认", "待指派"],
        ["TBD-04", "真实系统、人员、设备接口", "无法进入真实联调", "阶段一全部模拟", "待指派"],
    ], [1.8, 4.6, 3.7, 5.0, 3.0], 7.6)

    # 17
    add_heading(doc, "17 附录", 1)
    add_heading(doc, "17.1 需求追踪矩阵", 2)
    add_table(doc, ["需求", "功能", "数据对象", "接口", "验收"], [
        ["R-01", "FP-01", "Event/EventFact", "IF-01/IF-02", "TC-01/TC-02"],
        ["R-02", "FP-02", "Event/Device/MapFocus", "IF-07", "TC-01/TC-08"],
        ["R-03/R-04", "FP-03/FP-04", "ReasoningTrace/Knowledge/Rule/Graph", "IF-02/IF-03", "TC-03"],
        ["R-05", "FP-05", "CandidateStrategy/PlanVersion", "IF-02/IF-03", "TC-04"],
        ["R-06", "FP-06", "PlanVersion/Measure/Confirmation", "内部状态", "TC-05/TC-16"],
        ["R-07", "FP-07", "Dispatch/DispatchTarget", "IF-04", "TC-06/TC-07"],
        ["R-08", "FP-08", "DeviceEffect/DeviceState", "IF-05", "TC-07～TC-09"],
        ["R-09", "FP-09", "Device/GridRange/Filter", "内部状态", "TC-10～TC-12"],
        ["R-10", "FP-10", "Narrative", "IF-06", "TC-13/TC-14"],
        ["R-11", "FP-11", "TrafficSample/CapacityResult", "IF-03", "TC-15"],
        ["R-12", "FP-12", "TwinScenario/TwinFrame", "内部状态", "演示检查"],
        ["R-13", "FP-13", "AuditLog/FinalReport", "IF-08", "TC-17"],
        ["R-14", "FP-14", "配置/案例/测试报告", "本地运行", "TC-18"],
    ], [2.3, 2.4, 5.6, 3.0, 3.2], 7.8)
    add_heading(doc, "17.2 智能决策底座说明", 2)
    add_para(doc, "系统以“大模型+知识图谱+知识库+业务规则库”为智能决策底座。大模型负责事件语义理解、信息归纳和策略表达；知识图谱关联事件、道路、设施、资源、风险与措施，支持关联查询和依赖传播；知识库沉淀应急预案、历史案例、行业规范和处置经验，为策略提供可引用依据；业务规则库校验措施适用条件、资源冲突、执行边界及安全约束，对不符合强约束的措施实施阻断。")
    add_para(doc, "底座与交通流计算协同工作。交通流模型围绕瓶颈通行能力、行驶密度、排队密度和队尾演化进行计算，并保留输入、假设、公式口径和结果。系统据此生成多套候选管控策略，以统一结构呈现适用条件、实施措施、资源需求、预期效果、风险和置信依据。指挥员可横向比较方案，展开查看证据和规则校验，并沿事件事实—知识条目—图谱关系—规则—计算—策略—措施—回执链路追溯。由此实现策略可比较、研判过程可解释、决策依据可追溯。")
    add_heading(doc, "17.3 三流一致性审查报告", 2)
    add_para(doc, "审查范围覆盖第5章业务流程、第7章状态模型、第9章数据读写及第10章接口。审查方法为逐项建立“业务动作×状态变化×数据变化×接口/服务×闭环证据”矩阵，并检查异常分支与终态。")
    add_table(doc, ["业务动作", "状态变化", "数据读写", "接口/服务", "闭环证据"], [
        ["事件上报/加载", "draft→analyzing", "写Event/EventFact/Audit", "IF-01/IF-02", "事件卡片、结构化事实"],
        ["重复识别/融合", "incoming→merged或新建", "更新主事件/关联源事件", "事件中心", "合并理由、主事件ID"],
        ["智能研判", "trace pending→generated", "写Trace/Evidence/RuleCheck/TrafficCalc", "IF-02/IF-03", "五步推理与证据引用"],
        ["生成/比较策略", "generated→compared→selected", "写CandidateStrategy/PlanVersion", "策略中心", "统一字段与选定记录"],
        ["措施确认", "pending→confirmed/rejected", "写Confirmation/MeasureState", "时序引擎", "确认人、时间、版本"],
        ["模拟下发", "pending→dispatching→终态", "写Dispatch/Target/Receipt", "IF-04", "摘要、耗时、失败原因"],
        ["设备同步", "unchanged→synced/sync_failed", "写DeviceEffect/MapMarkerState", "IF-05", "网格/GIS一致、措施可反查"],
        ["态势与能力监测", "generating→success/fallback；sampling→等级", "写Narrative/TrafficSample/Capacity", "IF-03/IF-06", "30秒讲解、趋势与降级来源"],
        ["事实变化与版本替换", "active→superseded；confirmed→obsolete", "新PlanVersion并保留旧版", "TMS依赖传播", "版本链与撤销原因"],
        ["结案与导出", "active→closing→closed", "写FinalReport/Audit/Export", "IF-08", "报告、JSONL、审计"],
    ], [3.3, 3.6, 5.0, 2.7, 3.4], 7.2)
    add_heading(doc, "17.3.1 审查结论", 3)
    add_table(doc, ["审查项", "结论", "说明"], [
        ["业务流闭环", "通过", "从事件接入到结案归档完整；事实变化可回到策略生成形成修订环。"],
        ["状态流闭环", "通过", "事件、方案、措施、下发、设备、简报均定义初态、流转、异常与终态。"],
        ["数据流闭环", "通过", "核心动作均明确读写对象；策略、回执、设备变化和审计可关联。"],
        ["异常流闭环", "通过", "模型、地图、下发和交通样本均有可见降级或失败处理。"],
        ["跨流一致性", "通过", "业务动作、状态变化、数据变化和接口逐项对应，无孤立功能点。"],
        ["生产化完备性", "待后续确认", "负责人、真实接口、SLA/RTO/RPO、等保、容量和安全运营未纳入阶段一。"],
    ], [4.0, 3.0, 10.0])
    add_heading(doc, "17.4 阶段五确认记录", 2)
    add_table(doc, ["确认项", "结论"], [
        ["建设方式", "分阶段建设；当前功能作为阶段一基线。"],
        ["用户角色", "阶段一仅值班指挥员。"],
        ["执行边界", "分级执行；低风险自动模拟，高风险人工确认。"],
        ["智能底座", "大模型、知识图谱、知识库、业务规则库均为阶段一核心并纳入验收。"],
        ["数据与接口", "全部采用模拟数据，不接真实业务系统和物理设备。"],
        ["部署与显示", "Windows单机、本地浏览器、1920×1080、100%缩放。"],
        ["负责人", "产品、技术、业务、数据安全负责人待指派。"],
        ["生产指标", "SLA、RTO、RPO、等保等后续阶段确认。"],
        ["版本", "初始版本V0.1，由用户决定是否定版升级。"],
    ], [5.0, 11.0])

    # Final statement
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    r = p.add_run("— 文档结束 —")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("V0.1 评审稿｜所有未定生产化事项均须在后续阶段完成确认").font.size = Pt(10)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "路网智能管控需求规格说明书"
    doc.core_properties.subject = "阶段一演示基线 PRD"
    doc.core_properties.author = "Codex（待项目组复核）"
    doc.core_properties.keywords = "路网智能管控, PRD, GIS, 智能决策, 交通应急"
    doc.save(str(output))


def validate_document(path: Path) -> None:
    doc = Document(str(path))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    joined = "\n".join(texts)
    required = [
        "1 项目概述", "2 用户与使用场景", "3 需求总览", "4 信息架构与页面结构",
        "5 业务流程与交互时序", "6 功能模块设计", "7 详细功能规格", "8 界面与交互设计",
        "9 数据设计", "10 接口与集成", "11 权限与安全", "12 埋点、指标与审计",
        "13 非功能与运行保障", "14 测试与验收", "15 发布、部署与培训",
        "16 风险与待确认事项", "17 附录",
    ]
    missing = [item for item in required if item not in joined]
    if missing:
        raise RuntimeError(f"Missing required chapters: {missing}")
    if len(doc.tables) < 35:
        raise RuntimeError(f"Unexpectedly low table count: {len(doc.tables)}")
    if path.stat().st_size < 35000:
        raise RuntimeError(f"Output file is unexpectedly small: {path.stat().st_size}")
    print(f"VALID: {path}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} bytes={path.stat().st_size}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    template = Path(args.template).resolve()
    output = Path(args.output).resolve()
    if not template.exists():
        raise FileNotFoundError(template)
    build_document(template, output)
    validate_document(output)


if __name__ == "__main__":
    main()
